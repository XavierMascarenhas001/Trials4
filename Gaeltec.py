# dashboard_mapped.py
import streamlit as st
import pandas as pd
import plotly.express as px
import re
import geopandas as gpd
import pydeck as pdk
import os
import glob
from PIL import Image
from io import BytesIO
import base64
from streamlit_plotly_events import plotly_events
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import requests
from streamlit import cache_data
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from collections import OrderedDict
from openpyxl.styles import Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.styles import Border, Side
import io
from io import BytesIO
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import numbers
import plotly.graph_objects as go
import numpy as np
# ---- Formatting & logos (after all sheets written) ----

# --- Page config for wide layout ---
st.set_page_config(
    page_title="Gaeltec Dashboard",
    layout="wide",  # <-- makes the dashboard wider
    initial_sidebar_state="expanded"
)

def prepare_dataframe(df):
    df = df.copy()
    df.columns = df.columns.str.strip().str.lower()

    if 'datetouse' in df.columns:
        df['datetouse_dt'] = pd.to_datetime(
            df['datetouse'],
            errors='coerce',
            dayfirst=True
        ).dt.normalize()

        # ✅ ADD THESE (CRITICAL FIX)
        df['datetouse_date'] = df['datetouse_dt'].dt.date
        df['datetouse_display'] = df['datetouse_dt'].dt.strftime("%d/%m/%Y")

    else:
        df['datetouse_dt'] = pd.NaT
        df['datetouse_date'] = None
        df['datetouse_display'] = "Missing"

    # Numeric columns
    for col in ['total', 'orig']:
        if col in df.columns:
            df[col] = (
                df[col].astype(str)
                .str.replace(" ", "")
                .str.replace(",", ".", regex=False)
                .astype(float)
            )

    return df

def sanitize_sheet_name(name: str) -> str:
    """
    Remove or replace invalid characters for Excel sheet names.
    Excel sheet names cannot contain: : \ / ? * [ ]
    """
    name = str(name)
    name = re.sub(r'[:\\/*?\[\]]', '_', name)
    name = re.sub(r'[^\x00-\x7F]', '_', name)
    return name[:31]


def poles_to_word(df: pd.DataFrame) -> BytesIO:
    doc = Document()

    # Defensive cleaning
    df = df.copy()
    df = df.replace(
        to_replace=["nan", "NaN", "None", None],
        value=""
    )

    grouped = df.groupby('pole', sort=False)

    for pole, group in grouped:
        pole_str = str(pole).strip()
        if not pole_str:
            continue

        # Ordered set using dict keys (preserves order, removes duplicates)
        unique_texts = OrderedDict()

        for _, row in group.iterrows():
            parts = []

            wi = str(row.get('Work instructions', '')).strip()
            comment = str(row.get('comment', '')).strip()

            if wi:
                parts.append(wi)

            if comment:
                parts.append(f"({comment})")

            if parts:
                text = " ".join(parts)

                # Normalize for deduplication
                normalized = text.lower().strip()

                unique_texts[normalized] = text

        if not unique_texts:
            continue

        # Bullet paragraph
        p = doc.add_paragraph(style='List Bullet')

        run_number = p.add_run(f"{pole_str} – ")
        run_number.bold = True
        run_number.font.name = 'Times New Roman'
        run_number.font.size = Pt(12)

        texts = list(unique_texts.values())

        for i, text in enumerate(texts):
            run_item = p.add_run(text)
            run_item.bold = True
            run_item.font.name = 'Times New Roman'
            run_item.font.size = Pt(12)

            if "Erect Pole" in text:
                run_item.font.highlight_color = WD_COLOR_INDEX.RED

            if i < len(texts) - 1:
                p.add_run(" ; ")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def build_export_df(filtered_df):
    export_df = filtered_df.copy()

    # Rename columns
    export_df = export_df.rename(columns=column_rename_map)

    # Keep only columns that actually exist
    existing_cols = [c for c in export_columns if c in export_df.columns]
    export_df = export_df[existing_cols]

    return export_df

# Normalize strings: remove leading/trailing spaces, lowercase, remove extra dots
def normalize_item(s):
    if pd.isna(s):
        return ""
    s = str(s).strip().lower()           # strip spaces and lowercase
    s = s.replace(".", "")               # remove dots
    s = re.sub(r"\s+", " ", s)          # collapse multiple spaces
    return s

def apply_common_filters(df):
    df = df.copy()

    # Ensure datetime
    df['datetouse_dt'] = pd.to_datetime(
        df['datetouse'],
        errors='coerce',
        dayfirst=True
    )

    df['datetouse_dt'] = df['datetouse_dt'].dt.normalize()

    # ✅ ADD THESE BACK
    df['datetouse_date'] = df['datetouse_dt'].dt.date
    df['datetouse_display'] = df['datetouse_dt'].dt.strftime("%d/%m/%Y")

    # Date rule: after 2023
    df = df[df['datetouse_dt'].dt.year > 2023]

    # Segment filter
    if selected_segment != 'All' and 'segmentcode' in df.columns:
        df = df[
            df['segmentcode'].astype(str).str.strip()
            == str(selected_segment).strip()
        ]

    # Pole filter
    if selected_pole != "All" and 'pole' in df.columns:
        df = df[
            df['pole'].astype(str).str.strip()
            == str(selected_pole).strip()
        ]

    # Numeric
    if 'total' in df.columns:
        df['total'] = pd.to_numeric(df['total'], errors='coerce')

    return df.dropna(subset=['datetouse_dt'])
    
def prepare_dataframe(df):
    df = df.copy()
    df.columns = df.columns.str.strip().str.lower()

    if 'datetouse' in df.columns:
        df['datetouse_dt'] = pd.to_datetime(df['datetouse'], errors='coerce').dt.normalize()
    else:
        df['datetouse_dt'] = pd.NaT

    # Make numeric columns safe
    for col in ['total', 'orig']:
        if col in df.columns:
            df[col] = (
                df[col].astype(str)
                .str.replace(" ", "")
                .str.replace(",", ".", regex=False)
                .astype(float)
            )

    return df

def multi_select_filter(col, label, df):
    if col not in df.columns:
        return ["All"], df

    options = ["All"] + sorted(df[col].dropna().astype(str).unique())
    selected = st.sidebar.multiselect(label, options, default=["All"])

    if "All" in selected:
        return selected, df

    return selected, df[df[col].astype(str).isin(selected)]

def preprocess_df(df):
    import pandas as pd

    # Ensure the 'datetouse' column exists
    if 'datetouse' in df.columns:
        df['datetouse_dt'] = pd.to_datetime(df['datetouse'], errors='coerce').dt.normalize()
        df['datetouse_display'] = df['datetouse_dt'].dt.strftime("%d/%m/%Y")
    else:
        df['datetouse_dt'] = pd.NaT
        df['datetouse_display'] = "Missing"

    # Add other preprocessing here if needed
    # e.g., stripping strings, renaming columns, converting numbers
    return df

def to_excel(project_df, team_df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:

        # ---- Sheet 1: Revenue per Project ----
        if not project_df.empty:
            project_df.to_excel(writer, index=False, sheet_name="Revenue per Project", startrow=1)
            ws_proj = writer.sheets["Revenue per Project"]

            # ---- Column widths ----
            ws_proj.column_dimensions["A"].width = 30
            ws_proj.column_dimensions["B"].width = 18

            # ---- Styles ----
            header_font = Font(bold=True, size=14)
            header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
            thin_side = Side(style="thin")
            medium_side = Side(style="medium")
            thick_side = Side(style="thick")
            light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
            white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

            max_col = ws_proj.max_column
            max_row = ws_proj.max_row

            # Set row 1 height for images
            ws_proj.row_dimensions[1].height = 120

            # Header → row 2
            for col_idx, cell in enumerate(ws_proj[2], start=1):
                cell.font = header_font
                cell.fill = header_fill
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

            # Data rows → start row 3
            for row_idx in range(3, max_row + 1):
                fill = light_grey_fill if row_idx % 2 == 0 else white_fill
                for col_idx in range(1, max_col + 1):
                    cell = ws_proj.cell(row=row_idx, column=col_idx)
                    cell.fill = fill
                    cell.border = Border(
                        left=thin_side,
                        right=thin_side,
                        top=thin_side,
                        bottom=thin_side
                    )

            # ---- Add images in row 1 ----
            img1 = XLImage("Images/GaeltecImage.png")
            img2 = XLImage("Images/SPEN.png")
            img1.width = 120; img1.height = 120; img1.anchor = "A1"
            img2.width = 360; img2.height = 120; img2.anchor = "B1"
            ws_proj.add_image(img1)
            ws_proj.add_image(img2)

        # ---- Sheet 2: Revenue per Team ----
        if not team_df.empty:
            team_df.to_excel(writer, index=False, sheet_name="Revenue per Team", startrow=1)
            ws_team = writer.sheets["Revenue per Team"]

            ws_team.column_dimensions["A"].width = 25
            ws_team.column_dimensions["B"].width = 18

            max_col = ws_team.max_column
            max_row = ws_team.max_row

            # Row 1 for images
            ws_team.row_dimensions[1].height = 120

            # Header → row 2
            for col_idx, cell in enumerate(ws_team[2], start=1):
                cell.font = header_font
                cell.fill = header_fill
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

            # Data rows → start row 3
            for row_idx in range(3, max_row + 1):
                fill = light_grey_fill if row_idx % 2 == 0 else white_fill
                for col_idx in range(1, max_col + 1):
                    cell = ws_team.cell(row=row_idx, column=col_idx)
                    cell.fill = fill
                    cell.border = Border(
                        left=thin_side,
                        right=thin_side,
                        top=thin_side,
                        bottom=thin_side
                    )

            # ---- Add images in row 1 ----
            img1 = XLImage("Images/GaeltecImage.png")
            img2 = XLImage("Images/SPEN.png")
            img1.width = 120; img1.height = 120; img1.anchor = "A1"
            img2.width = 360; img2.height = 120; img2.anchor = "B1"
            ws_team.add_image(img1)
            ws_team.add_image(img2)

    output.seek(0)
    return output



def generate_excel_styled_multilevel(filtered_df, poles_df=None):
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    ws = wb.active
    ws.title = "Daily Revenue"

    # ---- Sheet 1: Daily Revenue ----
    if {'shire', 'project','region','segmentdesc', 'segmentcode', 'projectmanager', 'datetouse_dt','done', 'total','sourcefile'}.issubset(filtered_df.columns):
        daily_df = (
            filtered_df
            .groupby(['datetouse_dt','shire','project','region','segmentdesc','segmentcode','projectmanager','sourcefile'], as_index=False)
            .agg({'total':'sum'})
        )
        daily_df.rename(columns={
            'datetouse_dt':'Date',
            'total':'Revenue (£)',
            'region':'location',
            'segmentdesc':'Detail',
            'segmentcode':'Segment',
            'projectmanager':'Project Manager',
            'sourcefile':'Control file'
        }, inplace=True)

        # Write header in ROW 2 (row 1 reserved for images)
        for col_idx, col_name in enumerate(daily_df.columns.tolist(), start=1):
            ws.cell(row=2, column=col_idx, value=col_name)

        # Write data starting from row 3
        for r_idx, row in enumerate(daily_df.values.tolist(), start=3):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(row=r_idx, column=c_idx, value=value)

    # ---- Sheet 2: Poles Summary ----
    ws_summary = wb.create_sheet(title="Poles Summary")
    if poles_df is not None and not poles_df.empty:
        poles_summary = (
            poles_df[['shire','project','segmentcode','pole']]
            .drop_duplicates()
            .groupby(['shire','project','segmentcode'], as_index=False)
            .agg({'pole': lambda x: ', '.join(sorted(x.astype(str)))})
        )
        poles_summary.rename(columns={'pole':'Poles', 'segmentcode':'Segment'}, inplace=True)

        # Write multi-level headers (Row 2-4)
        headers = ['Shire','Project','Segment','location_map','Poles']
        for idx, h in enumerate(headers, start=1):
            ws_summary.cell(row=2, column=idx, value=h)  # Shire header
            ws_summary.cell(row=3, column=idx, value=h if h != 'Poles' else '')  # Project header
            ws_summary.cell(row=4, column=idx, value=h if h != 'Poles' else '')  # Segment header

        # Write data starting from row 5
        for r_idx, row in enumerate(poles_summary.values.tolist(), start=5):
            for c_idx, value in enumerate(row, start=1):
                ws_summary.cell(row=r_idx, column=c_idx, value=value)

    # ---- Formatting styles ----
    header_font = Font(bold=True, size=16)
    header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
    thin_side = Side(style="thin")
    medium_side = Side(style="medium")
    thick_side = Side(style="thick")
    light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    # ---- Add images ----
    IMG_HEIGHT = 120
    IMG_WIDTH_SMALL = 120
    IMG_WIDTH_LARGE = IMG_WIDTH_SMALL * 3

    # Set row 1 height to fit images
    ws.row_dimensions[1].height = IMG_HEIGHT * 0.75  # approximate pixels → Excel points
    ws_summary.row_dimensions[1].height = IMG_HEIGHT * 0.75
    img1 = XLImage("Images/GaeltecImage.png")
    img2 = XLImage("Images/SPEN.png")

    # Position images (row 1)
    img1.anchor = "B1"
    img2.anchor = "A1"

    ws.add_image(img1)
    ws.add_image(img2)

    # Same for Summary
    img1_s = XLImage("Images/GaeltecImage.png")
    img2_s = XLImage("Images/SPEN.png")

    img1_s.width = IMG_WIDTH_SMALL
    img1_s.height = IMG_HEIGHT
    img1_s.anchor = "A1"

    img2_s.width = IMG_WIDTH_LARGE
    img2_s.height = IMG_HEIGHT
    img2_s.anchor = "B1"

    # Sheet 2 images
    img1_s = XLImage("Images/GaeltecImage.png")
    img2_s = XLImage("Images/SPEN.png")
    img1_s.width = IMG_WIDTH_SMALL; img1_s.height = IMG_HEIGHT; img1_s.anchor = "A1"
    img2_s.width = IMG_WIDTH_LARGE; img2_s.height = IMG_HEIGHT; img2_s.anchor = "B1"
    ws_summary.add_image(img1_s)
    ws_summary.add_image(img2_s)

    # ---- Apply formatting ----
    for sheet in [ws, ws_summary]:
        max_col = sheet.max_column
        max_row = sheet.max_row

        # Header rows
        for row_idx in range(2, 5 if sheet == ws_summary else 3):
            for col_idx in range(1, max_col + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.font = header_font
                cell.fill = header_fill
                sheet.column_dimensions[get_column_letter(col_idx)].width = 60 if col_idx == 1 else 20
                cell.border = Border(
                    left=thick_side if col_idx == 1 else medium_side,
                    right=thick_side if col_idx == max_col else medium_side,
                    top=thick_side,
                    bottom=thick_side
                )

        # DATA ROWS → after headers
        start_data_row = 5 if sheet == ws_summary else 3
        for row_idx in range(start_data_row, max_row + 1):
            fill = light_grey_fill if row_idx % 2 == 1 else white_fill
            for col_idx in range(1, max_col + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.fill = fill
                cell.border = Border(
                    left=thin_side,
                    right=thin_side,
                    top=thin_side,
                    bottom=thin_side
                )

    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output
   
# --- MAPPINGS ---

# --- Project Manager Mapping ---
project_mapping = {
    "Jonathon Mcclung": ["Ayrshire", "PCB"],
    "Gary MacDonald": ["Ayrshire", "LV"],
    "Jim Gaffney": ["Lanark", "PCB"],
    "Calum Thomson": ["Ayrshire", "Connections"],
    "Calum Thomsom": ["Ayrshire", "Connections"],
    "Calum Thompson": ["Ayrshire", "Connections"],
    "Andrew Galt": ["Ayrshire", "-"],
    "Henry Gordon": ["Ayrshire", "-"],
    "Jonathan Douglas": ["Ayrshire", "11 kV"],
    "Jonathon Douglas": ["Ayrshire", "11 kV"],
    "Matt": ["Lanark", ""],
    "Lee Fraser": ["Ayrshire", "Connections"],
    "Lee Frazer": ["Ayrshire", "Connections"],
    "Mark": ["Lanark", "Connections"],
    "Mark Nicholls": ["Ayrshire", "Connections"],
    "Cameron Fleming": ["Lanark", "Connections"],
    "Ronnie Goodwin": ["Lanark", "Connections"],
    "Ian Young": ["Ayrshire", "Connections"],
    "Matthew Watson": ["Lanark", "Connections"],
    "Aileen Brese": ["Ayrshire", "Connections"],
    "Mark McGoldrick": ["Lanark", "Connections"]
}

# --- Region Mapping ---
mapping_region = {
    "Newmilns": ["Irvine Valley"],
    "New Cumnock": ["New Cumnock"],
    "Kilwinning": ["Kilwinning"],
    "Stewarton": ["Irvine Valley"],
    "Kilbirnie": ["Kilbirnie and Beith"],
    "Coylton": ["Ayr East"],
    "Irvine": ["Irvine Valley", "Irvine East", "Irvine West"],
    "TROON": ["Troon"],
    "Ayr": ["Ayr East", "Ayr North", "Ayr West"],
    "Maybole": ["Maybole, North Carrick and Coylton"],
    "Clerkland": ["Irvine Valley"],
    "Glengarnock": ["Kilbirnie and Beith"],
    "Ayrshire": ["North Coast and Cumbraes","Prestwick", "Saltcoats and Stevenston", "Troon", "Ayr East", "Ayr North",
                 "Ayr West","Annick","Ardrossan and Arran","Dalry and West Kilbride","Girvan and South Carrick","Irvine East",
                 "Irvine Valley","Irvine West","Kilbirnie and Beith","Kilmarnock East and Hurlford","Kilmarnock North",
                 "Kilmarnock South","Kilmarnock West and Crosshouse","Kilwinning","Kyle","Maybole, North Carrick and Coylton",
                 "Ayr, Carrick and Cumnock","East_Ayrshire","North_Ayrshre","South_Ayrshre","Doon Valley"],
    "Lanark": ["Abronhill, Kildrum and the Village","Airdrie Central","Airdrie North","Airdrie South","Avondale and Stonehouse",
               "Ballochmyle","Bellshill","Blantyre","Bothwell and Uddingston","Cambuslang East","Cambuslang West",
               "Clydesdale East","Clydesdale North","Clydesdale South","Clydesdale West","Coatbridge North and Glenboig",
               "Coatbridge South","Coatbridge West","Cumbernauld North","Cumbernauld South",
               "East Kilbride Central North","East Kilbride Central South","East Kilbride East","East Kilbride South",
               "East Kilbride West","Fortissat","Hamilton North and East","Hamilton South","Hamilton West and Earnock",
               "Mossend and Holytown","Motherwell North","Motherwell South East and Ravenscraig","Motherwell West",
               "Rutherglen Central and North","Rutherglen South","Strathkelvin","Thorniewood","Wishaw","Larkhall",
               "Airdrie and Shotts","Cumbernauld, Kilsyth and Kirkintilloch East","East Kilbride, Strathaven and Lesmahagow",
               "Lanark and Hamilton East","Motherwell and Wishaw","North_Lanarkshire","South_Lanarkshire"]
}

# --- File Project Mapping ---
file_project_mapping = {
    "pcb 2022": ["Ayrshire", "PCB"],
    "33kv refurb": ["Ayrshire", "33kv Refurb"],
    "connections": ["Ayrshire", "Connections"],
    "storms": ["Ayrshire", "Storms"],
    "11kv refurb": ["Ayrshire", "11kv Refurb"],
    "11kV Refurb Ayrshire 2026": ["Ayrshire", "11kV Refurb"],
    "11kV Ref Ayr Pinwherry": ["Ayrshire", "11kV Refurb"],
    "aurs road": ["Ayrshire", "Aurs Road"],
    "spen labour": ["Ayrshire", "SPEN Labour"],
    "lvhi5": ["Ayrshire", "LV"],
    "pcb": ["Ayrshire", "PCB"],
    "lanark": ["Lanark", ""],
    "11kv refur": ["Lanark", "11kv Refurb"],
    "lv & esqcr": ["Lanark", "LV"],
    "11kv rebuilt": ["Lanark", "11kV Rebuilt"],
    "33kv rebuilt": ["Lanark", "33kV Rebuilt"],
    "Hi5_4_Lanark_2026": ["Lanark", "11kV Refurb"],
    "Hi5_4_Ayrshire_2026": ["Lanark", "11kV Refurb"],

}

CV7_erect = {
    "Erect Single HV/EHV Pole, up to and including 12 metre pole":"CV7 HV pole", 
    "Erect Single HV/EHV Pole, up to and including 12 metre pole.":"CV7  HV pole",
}

CV7_erect_H = {
    "Erect Section Structure 'H' HV/EHV Pole, up to and including 12 metre pole.":"CV7 HV pole"
}

CV7_erect_lv = {
    "Erect LV Structure Single Pole, up to and including 12 metre pole" :"CV7 LV pole",
}

CV7_recover = {
    "Recover single pole, up to and including 15 metres in height, and reinstate, all ground conditions":"CV7",
    "Recover 'A' / 'H' pole, up to and including 15 metres in height, and reinstate, all ground conditions":"CV7  HV pole"
}


# --- Transformer Mappings ---
CV7_Tx = {
    "Erect pole mounted transformer up to 100kVA 1.ph.": "CV7 Tx",
    "Erect pole mounted transformer up to 200kVA 3.p.h.": "CV7 Tx",
    "Erect Voltage Regulator.": "CV7 Tx",
    "Erect Voltage Transformer (VT), RTU or Repeater": "CV7 Tx",
    "Erect 12kV/36kV Surge arrestors ( directly mounted ).": "CV7 Tx)",
    "Remove pole mounted tranformer.": "CV7 Tx)",
    "Remove platform mounted or 'H' pole mounted transformer.": "CV7 Tx)"
}

# --- Transformer Mappings ---
transformer = {
    "Transformer 1ph 50kVA": "TX 1ph (50kVA)",
    "Transformer 3ph 50kVA": "TX 3ph (50kVA)",
    "Transformer 1ph 100kVA": "TX 1ph (100kVA)",
    "Transformer 1ph 25kVA": "TX 1ph (25kVA)",
    "Transformer 3ph 200kVA": "TX 3ph (200kVA)",
    "Transformer 3ph 100kVA": "TX 3ph (100kVA)"
}

# --- Equipment / Conductor Mappings ---
CV7_OHL_CONDUCTOR_instal = {
    "Install bare conductor, run out, sag, terminate, bind in and connect jumpers; <100mm²": "CV7 OHL CONDUCTOR",
    "Install bare conductor, run out, sag, terminate, bind in and connect jumpers; >=100mm² <200mm²": "CV7 OHL CONDUCTOR",
    "Install conductor, run out, sag, terminate, clamp in and form jumper loops; >=200mm²": "CV7 OHL CONDUCTOR",

}
CV7_OHL_CONDUCTOR_recover = {
    "Recover overhead wire and fittings; HV/EHV overhead line or Hardex Pilot (1 conductor)": "CV7 OHL CONDUCTOR",
    "Recover overhead wire and fittings; HV/EHV overhead line or Hardex Pilot (2 conductor)": "CV7 OHL CONDUCTOR",
    "Recover overhead wire and fittings; HV/EHV overhead line or Hardex Pilot (3 conductor)": "CV7 OHL CONDUCTOR",
}

    # LV cables per meter
CV7_OHL_CONDUCTOR_LV_instal = {
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 2c": "CV7 OHL CONDUCTOR LV",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 4c": "CV7 OHL CONDUCTOR LV",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 2c + Earth": "CV7 OHL CONDUCTOR LV",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 4c + Earth": "CV7 OHL CONDUCTOR LV",
}

CV7_OHL_CONDUCTOR_LV_recover = {
    "Recover overhead wires and fittings; LV openwire overhead line (2 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV openwire overhead line (3 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV openwire overhead line (4 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV openwire overhead line (5 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV service overhead line (open, concentric or ABC, 2 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV service overhead line (open, concentric or ABC, 3 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV service overhead line (open, concentric or ABC, 4 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover overhead wires and fittings; LV service overhead line (open, concentric or ABC, 5 conductors)": "CV7 OHL CONDUCTOR LV",
    "Recover cleated service": "CV7 OHL CONDUCTOR LV",
}


CV7_SWITCHGEAR = {
    "Erect 11kV/33kV ABSW": "CV7 SWITCHGEAR",
    "Erect 11kV Remote Controlled Switch Disconnector ( Soule Auguste ) or Auto Reclosure unit c/w VT, Aerial, RTU & umbilical cable.": "CV7 SWITCHGEAR",
    "Erect 1.ph fuse units at single tee off pole or in line pole.": "CV7 SWITCHGEAR",
    "Erect 3.ph fuse units at single tee off pole or in line pole.": "CV7 SWITCHGEAR",
    "Additional cost for fitting fuse outrigger bracket.": "CV7 SWITCHGEAR",
    "Remove 11kV/33kV ABSW": "CV7 SWITCHGEAR",
}

CV7_UG = {
    "Installation of cable only in trench dug by others; 11kV Cable 3 x 1 core.": "CV7 UG 11 kV",
    "Install cable in existing duct; 11kV Cable 3 x 1 core.": "CV7 UG 11 kV",
    "Installation of cable only in trench dug by others; 33kV Cable 3 x 1 core.": "CV7 UG 33 kV",
    "Install cable in existing duct; 33kV Cable 3 x 1 core.": "CV7 UG 33 kV",
    "Installation of cable only in trench dug by others; LV Cable Large or 11kV Cable 1 x 3 Core": "CV7 UG",
    "Install cable in existing duct; LV Cable Large or 11kV Cable 1 x 3 Core": "CV7 UG",
    "Installation of cable only in trench dug by others; LV Service, Small LV or Pilot Cable.": "CV7 UG LV Service",
    "Install cable in existing duct; LV Service, Small LV or Pilot Cable.": "CV7 UG LV Service",
}

CV7_CB = {
    "Remove Auto Reclosure.": "CV7 CB",
}

Switch = {
    "Noja": "Noja",
    "11kV PMSW (Soule)": "11kV PMSW (Soule)",
    "11kv ABSW Hookstick Standard": "11kv ABSW Hookstick Standard",
    "11kv ABSW Hookstick Spring loaded mech": "11kv ABSW Hookstick Spring loaded mech",
    "33kv ABSW Hookstick Dependant": "33kv ABSW Hookstick Dependant",
}

Fuses = {
    "100A LV Fuse JPU 82.5mm": "100A LV Fuse JPU 82.5mm",
    "160A LV Fuse JPU 82.5mm": "160A LV Fuse JPU 82.5mm",
    "200A LV Fuse JPU 82.5mm": "200A LV Fuse JPU 82.5mm",
    "315A LV Fuse JPU 82.5mm": "315A LV Fuse JPU 82.5mm",
    "400A LV Fuse JPU 82.5mm": "400A LV Fuse JPU 82.5mm",
    "200A LV Fuse JSU 92mm": "200A LV Fuse JSU 92mm",
    "315A LV Fuse JSU 92mm": "315A LV Fuse JSU 92mm",
    "400A LV Fuse JSU 92mm": "400A LV Fuse JSU 92mm",
    "100A LV Fuse - Porcelain screw-in": "100A LV Fuse - Porcelain screw-in",
    "160A LV Fuse - Porcelain screw-in": "160A LV Fuse - Porcelain screw-in",
    "200A LV Fuse - Porcelain screw-in": "200A LV Fuse - Porcelain screw-in",
    "Single Phase cut out kit 100A Henley Series 7": "Single Phase cut out kit 100A Henley Series 7",
    "Three Phase cut out kit 100A Henley Series 7": "Three Phase cut out kit 100A Henley Series 7",
    "Three Phase 200A Cut out": "Three Phase 200A Cut out",
    "Cut out Fuse (MF) 60A": "Cut out Fuse (MF) 60A",
    "Cut out Fuse (MF) 80A": "Cut out Fuse (MF) 80A",
    "Cut out Fuse (MF) 100A": "Cut out Fuse (MF) 100A",
    "11KV FUSE UNIT - C-TYPE": "11KV FUSE UNIT - C-TYPE",
    "11KV SOLID LINK - C-TYPE": "11KV SOLID LINK - C-TYPE",
    "11KV OHL ASL C-TYPE RESET 20A 2 SHOT": "11KV OHL ASL C-TYPE RESET 20A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 25A 2 SHOT": "11KV OHL ASL C-TYPE RESET 25A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 40A 1 SHOT": "11KV OHL ASL C-TYPE RESET 40A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 40A 2 SHOT": "11KV OHL ASL C-TYPE RESET 40A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 1 SHOT": "11KV OHL ASL C-TYPE RESET 63A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 2 SHOT": "11KV OHL ASL C-TYPE RESET 63A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 3 SHOT": "11KV OHL ASL C-TYPE RESET 63A 3 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 1 SHOT": "11KV OHL ASL C-TYPE RESET 100A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 2 SHOT": "11KV OHL ASL C-TYPE RESET 100A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 3 SHOT": "11KV OHL ASL C-TYPE RESET 100A 3 SHOT",
    "11KV OHL FUSE ELEMENT C-TYPE 15A": "11KV OHL FUSE ELEMENT C-TYPE 15A",
    "11KV OHL FUSE ELEMENT C-TYPE 25A": "11KV OHL FUSE ELEMENT C-TYPE 25A",
    "11KV OHL FUSE ELEMENT C-TYPE 30A": "11KV OHL FUSE ELEMENT C-TYPE 30A",
    "11KV OHL FUSE ELEMENT C-TYPE 40A": "11KV OHL FUSE ELEMENT C-TYPE 40A",
    "11KV OHL FUSE ELEMENT C-TYPE 50A": "11KV OHL FUSE ELEMENT C-TYPE 50A",
    "11KV OHL ASL DJP-TYPE 20A 2 SHOT": "11KV OHL ASL DJP-TYPE 20A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 25A 1 SHOT": "11KV OHL ASL DJP-TYPE 25A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 25A 2 SHOT": "11KV OHL ASL DJP-TYPE 25A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 40A 1 SHOT": "11KV OHL ASL DJP-TYPE 40A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 40A 2 SHOT": "11KV OHL ASL DJP-TYPE 40A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 1 SHOT": "11KV OHL ASL DJP-TYPE 63A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 2 SHOT": "11KV OHL ASL DJP-TYPE 63A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 3 SHOT": "11KV OHL ASL DJP-TYPE 63A 3 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 1 SHOT": "11KV OHL ASL DJP-TYPE 100A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 2 SHOT": "11KV OHL ASL DJP-TYPE 100A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 3 SHOT": "11KV OHL ASL DJP-TYPE 100A 3 SHOT",
    "11KV OHL FUSE ELEMENT DJP-TYPE 15A": "11KV OHL FUSE ELEMENT DJP-TYPE 15A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 25A": "11KV OHL FUSE ELEMENT DJP-TYPE 25A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 30A": "11KV OHL FUSE ELEMENT DJP-TYPE 30A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 40A": "11KV OHL FUSE ELEMENT DJP-TYPE 40A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 50A": "11KV OHL FUSE ELEMENT DJP-TYPE 50A",
}


CV31 = {
    "Replace / Fit safety or warning sign, number plates or name plate": "CV31",
    "Barbed Wire Wrap ACD (or Enhanced) single pole or stay - Replace/Repair": "CV31",
    "Steelwork bonding repair / fit.": "CV31",
    "Replace LV/HV/Earth guard missing / damaged.": "CV31",
}


CV8 = {
    "Tighten existing stay.": "CV8",
    "Plumb single pole.": "CV8",
    "Erect/Replace stay above ground only.": "CV8",
    "Erect/Replace stay complete including block or driven type anchor": "CV8",
    "Erect/Replace stay complete including rock type anchor": "CV8",
    "Retrofit structure with Anchor Clamp fitting for Section / Angle / Terminal support": "CV8",
    "Erect Single Crossarm to single pole.": "CV8",
    "Erect Double Crossarm to single pole": "CV8",
    "Erect Double Crossarm 'H' Pole formation": "CV8",
    "Remove Steelwork crossarm item only": "CV8",
    "Change 11kV Insulators to avoid contamination from old conductor": "CV8",
    "Change 33kV Insulators to avoid contamination from old conductor": "CV8",
    "Replace tension insulator, 11kV.": "CV8",
    "Replace tension insulator, 33kV.": "CV8",
    "Replace / Fit high visibility stay guard": "CV8",
    "Additional cost for fitting Stay Outrigger Bracket": "CV8",
    "Additional cost for fitting Angle / Terminal stay attachment plates on Heavy Construction as SP4009862": "CV8",
    "Recover and reinstate stay position,all ground conditions.": "CV8",
    "Fit foundation block to existing pole.": "CV8",
    "Fit bog shoe foundation to existing single pole.": "CV8",
    "Replace jumper / dropper mechanical connection with compression connection": "CV8",
    "Replace jumper / dropper with live line bail and flexible jumper conductor": "CV8",
    "Replace / Repair conductor with mid span joint using compression connection": "CV8",
    "Conductor repair; piece in conductor including compression joints": "CV8",
    "Bind In Conductors; 1.ph 11kV Intermediate / Pin Angle pole.": "CV8",
    "Bind In Conductors; 3.ph 11kV Intermediate / Pin Angle pole.": "CV8",
    "Conductor Terminations - 1.ph 11kV Section pole including jumpers.": "CV8",
    "Conductor Terminations - 3.ph 11kV Section pole including jumpers.": "CV8",
    "Conductor Terminations - 1.ph 11kV Terminal pole.": "CV8",
    "Conductor Terminations - 3.ph 11kV Terminal pole.": "CV8",
    "Unbind and reregulate existing conductors": "CV8",
    "Remove Steelwork crossarm item only": "CV8",
    "Convert 1.ph 11kV Intermediate pole into Section Pole.": "CV8",
    "Convert 1.ph/3.p.h. 11kV line pole into Terminal Pole.": "CV8",
    "Convert 3.ph 11kV Intermediate pole into Section Pole.": "CV8",
    "Change 11kV Insulators to avoid contamination from old conductor": "CV8",
    "Change 33kV Insulators to avoid contamination from old conductor": "CV8",
    "Replace 11kV/33kV insulator pin and insulator, including unbinding and binding in": "CV8",
    "Replace 11kV/33kV insulator binder": "CV8",
    "Replace tension insulator, 11kV": "CV8",
    "Replace tension insulator, 33kV": "CV8",
    "Replace 11kV/33kV dead end termination": "CV8",
    "Additional cost for erection of pilot pin and insulator or pilot post insulator (11kV or 33kV)": "CV8",
    "Replace insulated conductor HV/LV earth above ground to first rod": "CV8",
    "Install Copper Covered Green / Yellow HV Earth or Black LV Earth to foot of pole": "CV8",
    "Install EHV/ HV Earth Electrode including excavate & reinstate (up to 8mtrs)": "CV8",
    "Install LV Earth Electrode including excavate & reinstate (up to 28mtrs)": "CV8",
    "Additional extra over for additional earthing excavated, laid & backfilled": "CV8",
    "Install Earth Electrode within cable trench": "CV8",
    "Erect 11kV Cable Termination ( incorporating surge arrestors )": "CV8",
    "Erect 33kV Cable Termination ( incorporating surge arrestors )": "CV8",
    "Steelwork bonding repair / fit": "CV8",
    "Replace LV/HV/Earth guard missing / damaged": "CV8",
    "Erect 1.ph LV cable pole termination": "CV8",
    "Erect 3.ph LV cable pole termination": "CV8",
    "Remove 11kV/33kV Cable termination": "CV8",
    "Remove LV cable termination": "CV8",
}

summary_items = [
    "Erect Single HV/EHV Pole, up to and including 12 metre pole.",
    "Erect Section Structure 'H' HV/EHV Pole, up to and including 12 metre pole",
    "Erect LV Structure Single Pole, up to and including 12 metre pole",
    "Recover single pole, up to and including 15 metres in height, and reinstate, all ground conditions",
    "Recover 'A' / 'H' pole, up to and including 15 metres in height, and reinstate, all ground conditions",
    "Erect 11kV/33kV ABSW.",
    "Remove 11kV/33kV ABSW",
    "Noja"
    "0.5 kVa Tx for Noja"
    "11kV PMSW (Soule)"
    "Remove Auto Reclosure",
    "Erect pole mounted transformer up to 100kVA 1.ph",
    "Erect pole mounted transformer up to 200kVA 3.p.h",
    "Remove pole mounted transformer",
    "Remove platform mounted or 'H' pole mounted transformer",
    "Install bare conductor, run out, sag, terminate, bind in and connect jumpers; <100mm²",
    "Install bare conductor, run out, sag, terminate, bind in and connect jumpers; >=100mm² <200mm²",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 2c + Earth",
    "Install conductor, run out, sag, terminate, clamp in and connect jumpers; 4c + Earth",
    "Install service span including connection to mainline & building / structure",
    "Erect 3.ph fuse units at single tee off pole or in line pole"
    "Remove 1.ph or 3.ph HV fuses",    
]

categories = [
    ("CV7_erect", CV7_erect, "Quantity"),
    ("CV7_erect_H", CV7_erect_H, "Quantity"),
    ("CV7_erect_LV", CV7_erect_lv, "Quantity"),
    ("CV7_recover", CV7_recover, "Quantity"),
    ("CV7 Tx", CV7_Tx, "Quantity"),
    ("transformer", transformer, "Quantity"),
    ("CV7 OHL CONDUCTOR_instal", CV7_OHL_CONDUCTOR_instal, "Length (Km)"),
    ("CV7 OHL CONDUCTOR_recover", CV7_OHL_CONDUCTOR_recover, "Length (m)"),
    ("CV7 OHL CONDUCTOR LV_instal", CV7_OHL_CONDUCTOR_LV_instal, "Length (Km)"),
    ("CV7 OHL CONDUCTOR LV_recover", CV7_OHL_CONDUCTOR_LV_recover, "Length (m)"),
    ("CV7 SWITCHGEAR", CV7_SWITCHGEAR, "Quantity"),
    ("CV7_UG", CV7_UG, "Quantity"),
    ("CV7_CB", CV7_CB, "Quantity"),
    ("Switch", Switch, "Quantity"),
    ("Fuses", Fuses, "Quantity"),
    ("CV31", CV31, "Quantity"),
]

column_rename_map = {
    "mapped": "Output",
    "segmentcode": "Circuit",
    "datetouse_display": "Date",
    "qty": "Quantity_original",
    "qcvi":"qcvi",
    "qsub": "Quantity_used",
    "segmentdesc": "Segment",
    "shire": "District",
    "pid_ohl_nr": "PID",
    "projectmanager": "Project Manager"
}

export_columns = [
    'Output','comment', 'item', 'Quantity_original','qcvi','Quantity_used', 'pole', 'Date',
    'District', 'project', 'Project Manager', 'Circuit', 'Segment',
    'team lider', 'PID','total', 'orig', 'sourcefile'
]

# --- Gradient background ---
gradient_bg = """
<style>
    .stApp {
        background: linear-gradient(
            90deg,
            rgba(41, 28, 66, 1) 10%, 
            rgba(36, 57, 87, 1) 35%
        );
        color: white;
    }
</style>
"""
st.markdown(gradient_bg, unsafe_allow_html=True)

# --- Load logos ---
logo_left = Image.open(r"Images/GaeltecImage.png").resize((80, 80))
logo_right = Image.open(r"Images/SPEN.png").resize((160, 80))

# --- Header layout ---
col1, col2, col3 = st.columns([1, 4, 1])
with col1: st.image(logo_left)
with col2: st.markdown("<h1 style='text-align:center; margin:0;'>Gaeltec Utilities.UK</h1>", unsafe_allow_html=True)
with col3: st.image(logo_right)
st.markdown("<h1>📊 Data Management Dashboard</h1>", unsafe_allow_html=True)

# -------------------------------
# --- File Upload & Initial DF ---
# App Header
# -------------------------------
st.header("Upload Data Files")

# -------------------------------
# Load Aggregated Parquet
# -------------------------------
@st.cache_data
def load_master(file):
    df = pd.read_parquet(file, engine='pyarrow')  # pyarrow is faster
    df = preprocess_df(df)                        # preprocess once
    return df

master_file = st.file_uploader("Upload Master.parquet", type=["parquet"], key="master")
base_df = None

if master_file is not None:
    base_df = load_master(master_file)



# -------------------------------
# Date Source Selector
# -------------------------------
date_source = st.sidebar.radio(
    "Select Date Source",
    ["Planned + Done (datetouse)", "Done Only (done)"]
)

# -------------------------------
# --- Team Filter (GLOBAL) ---
# -------------------------------
base_df = None

if master_file:
    base_df = pd.read_parquet(master_file)
    base_df.columns = base_df.columns.str.strip().str.lower()

    # Normalize date
    if date_source == "Planned + Done (datetouse)":
        if 'datetouse' in base_df.columns:
            base_df['datetouse_dt'] = pd.to_datetime(base_df['datetouse'], errors='coerce').dt.normalize()
        else:
            base_df['datetouse_dt'] = pd.NaT
    elif date_source == "Done Only (done)":
        if 'done' in base_df.columns:
            base_df['datetouse_dt'] = pd.to_datetime(base_df['done'], errors='coerce').dt.normalize()
        else:
            base_df['datetouse_dt'] = pd.NaT

    # Normalize numeric columns
    for col in ['total', 'orig']:
        if col in base_df.columns:
            base_df[col] = (
                base_df[col]
                .astype(str)
                .str.replace(" ", "")
                .str.replace(",", ".", regex=False)
            )
            base_df[col] = pd.to_numeric(base_df[col], errors='coerce')

# Stop early if no data
if base_df is None:
    st.info("Please upload Master.parquet to continue.")
    st.stop()

# -------------------------------
# Sidebar Filters
# -------------------------------
st.sidebar.header("Filter Options")

def multiselect_filter(df, column, label):
    if column not in df.columns:
        return ["All"], df
    options = ["All"] + sorted(df[column].dropna().astype(str).unique())
    selected = st.sidebar.multiselect(label, options, default=["All"])
    if "All" not in selected:
        df = df[df[column].astype(str).isin(selected)]
    return selected, df

filtered_df = base_df.copy()

selected_shire, filtered_df = multiselect_filter(filtered_df, 'shire', "Select Shire")
selected_project, filtered_df = multiselect_filter(filtered_df, 'project', "Select Project")
selected_pm, filtered_df = multiselect_filter(filtered_df, 'projectmanager', "Select Project Manager")
selected_segment, filtered_df = multiselect_filter(filtered_df, 'segmentcode', "Select Segment Code")
selected_pole, filtered_df = multiselect_filter(filtered_df, 'pole', "Select Pole")
selected_type, filtered_df = multiselect_filter(filtered_df, 'type', "Select Type")
selected_team, filtered_df = multiselect_filter(filtered_df, 'team_name', "Select Team")


# -------------------------------
# Date Filter
# -------------------------------
# -------------------------------
# --- Standardize Date Column ---
# -------------------------------
# Convert any existing datetime column to just date (no hours/minutes/seconds)
# Use the already prepared base_df (which respects the radio button)
filtered_df = base_df.copy()

# Ensure datetouse_dt stays datetime (safety check)
filtered_df['datetouse_dt'] = pd.to_datetime(
    filtered_df['datetouse_dt'], errors='coerce'
).dt.normalize()

# Create display column ONLY (do not overwrite datetime)
filtered_df['datetouse_display'] = filtered_df['datetouse_dt'] \
    .dt.strftime("%d/%m/%Y") \
    .fillna("Missing")


filter_type = st.sidebar.selectbox(
    "Filter by Date",
    ["Single Day", "Week", "Month", "Year", "Custom Range", "Unplanned"]
)

date_range_str = ""
if filter_type == "Unplanned":
    filtered_df = filtered_df[filtered_df['datetouse_dt'].isna()]
    date_range_str = "Unplanned"
else:
    filtered_df = filtered_df[filtered_df['datetouse_dt'].notna()]

    if filter_type == "Single Day":
        d = st.sidebar.date_input("Select date")
        d_ts = pd.Timestamp(d)
        filtered_df = filtered_df[filtered_df['datetouse_dt'] == d_ts]
        date_range_str = d.strftime("%d/%m/%Y")

    elif filter_type == "Week":
        start = pd.Timestamp(st.sidebar.date_input("Week start"))
        end = start + pd.Timedelta(days=6)
        filtered_df = filtered_df[
            (filtered_df['datetouse_dt'] >= start) &
            (filtered_df['datetouse_dt'] <= end)
        ]
        date_range_str = f"{start.strftime('%d/%m/%Y')} → {end.strftime('%d/%m/%Y')}"

    elif filter_type == "Month":
        d = st.sidebar.date_input("Pick any date in month")
        filtered_df = filtered_df[
            (filtered_df['datetouse_dt'].dt.month == d.month) &
            (filtered_df['datetouse_dt'].dt.year == d.year)
        ]
        date_range_str = d.strftime("%B %Y")

    elif filter_type == "Year":
        y = st.sidebar.number_input("Year", 2000, 2100, 2025)
        filtered_df = filtered_df[filtered_df['datetouse_dt'].dt.year == y]
        date_range_str = str(y)

    elif filter_type == "Custom Range":
        start = pd.Timestamp(st.sidebar.date_input("Start date"))
        end = pd.Timestamp(st.sidebar.date_input("End date"))
        filtered_df = filtered_df[
            (filtered_df['datetouse_dt'] >= start) &
            (filtered_df['datetouse_dt'] <= end)
        ]
        date_range_str = f"{start.strftime('%d/%m/%Y')} → {end.strftime('%d/%m/%Y')}"
        
    # -------------------------------
    # --- Total & Variation Display ---
    # -------------------------------
    total_sum, variation_sum = 0, 0
    if 'total' in filtered_df.columns:
        total_series = pd.to_numeric(filtered_df['total'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
                                     errors='coerce')
        total_sum = total_series.sum(skipna=True)
        if 'orig' in filtered_df.columns:
            orig_series = pd.to_numeric(filtered_df['orig'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
                                        errors='coerce')
            variation_sum = (total_series - orig_series).sum(skipna=True)

    formatted_total = f"{total_sum:,.2f}".replace(",", " ").replace(".", ",")
    formatted_variation = f"{variation_sum:,.2f}".replace(",", " ").replace(".", ",")

    # Money logo
    money_logo_path = r"Images/Pound.png"
    money_logo = Image.open(money_logo_path).resize((40, 40))
    buffered = BytesIO()
    money_logo.save(buffered, format="PNG")
    money_logo_base64 = base64.b64encode(buffered.getvalue()).decode()

    # Display Total & Variation (Centered)
    st.markdown("<h2>Financial</h2>", unsafe_allow_html=True)
    st.markdown("<h3 style='text-align:center; color:white;'>Revenue</h3>", unsafe_allow_html=True)
    try:
        st.markdown(
            f"""
            <div style='display:flex; justify-content:center;'>
                <div style='display:flex; flex-direction:column; gap:4px;'>
                    <div style='display:flex; align-items:center; gap:10px;'>
                        <h2 style='color:#32CD32; margin:0; font-size:36px;'><b>Total:</b> {formatted_total}</h2>
                        <img src='data:image/png;base64,{money_logo_base64}' width='40' height='40'/>
                    </div>
                    <div style='display:flex; align-items:center; gap:8px;'>
                        <h2 style='color:#32CD32; font-size:25px; margin:0;'><b>Variation:</b> {formatted_variation}</h2>
                        <img src='data:image/png;base64,{money_logo_base64}' width='28' height='28'/>
                    </div>
                    <p style='text-align:center; font-size:14px; margin-top:4px;'>
                        ({date_range_str}, Shires: {selected_shire}, Projects: {selected_project}, PMs: {selected_pm})
                    </p>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    except Exception as e:
        st.warning(f"Could not display Total & Variation: {e}")
# -------------------------------
# Revenue Over Time
# -------------------------------

general_summary = pd.DataFrame(
    columns=["Description", "Total Quantity", "Comment"]
)
if not filtered_df.empty and 'datetouse_dt' in filtered_df.columns and 'total' in filtered_df.columns:
    # Aggregate revenue per date
    revenue_df = (
        filtered_df
        .dropna(subset=['datetouse_dt'])
        .groupby('datetouse_dt', as_index=False)['total']
        .sum()
        .sort_values('datetouse_dt')
    )

    # Ensure datetime column
    revenue_df['datetouse_dt'] = pd.to_datetime(revenue_df['datetouse_dt'])

    fig = go.Figure()

    # Scatter points (all data)
    fig.add_trace(go.Scattergl(
        x=revenue_df['datetouse_dt'],
        y=revenue_df['total'],
        mode='markers',
        marker=dict(size=8, color='#FFA500'),
        name='Revenue'
    ))

    # Dashed line connecting points
    fig.add_trace(go.Scatter(
        x=revenue_df['datetouse_dt'],
        y=revenue_df['total'],
        mode='lines',
        line=dict(dash='dash', color='#FFA500'),
        name='Trend'
    ))

    # Layout with horizontal gridlines
    fig.update_layout(
        height=500,
        xaxis_title="Date",
        yaxis_title="Revenue (£)",
        hovermode="x unified",
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='white'),
        xaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.1)'),
        yaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.2)', zeroline=False)
    )

    st.plotly_chart(fig, use_container_width=True)
else:
    st.info("No data for selected filters.")

prepared_df = filtered_df.copy() if filtered_df is not None else pd.DataFrame()
# -------------------------------
# --- Mapping Bar Charts + Drill-down + Excel Export ---
# -------------------------------
st.header("🪵 Materials")
convert_to_miles = st.checkbox("Convert Equipment/Conductor Length to Miles")

categories = [
    ("CV7_erect", CV7_erect, "Quantity"),
    ("CV7_erect_H", CV7_erect_H, "Quantity"),
    ("CV7_erect_lv", CV7_erect_lv, "Quantity"),
    ("CV7_recover", CV7_recover, "Quantity"),
    ("CV7 Tx", CV7_Tx, "Quantity"),
    ("transformer", transformer, "Quantity"),
    ("CV7 OHL CONDUCTOR_instal", CV7_OHL_CONDUCTOR_instal, "Length (Km)"),
    ("CV7 OHL CONDUCTOR_recover", CV7_OHL_CONDUCTOR_recover, "Length (m)"),
    ("CV7 OHL CONDUCTOR LV_instal", CV7_OHL_CONDUCTOR_LV_instal, "Length (Km)"),
    ("CV7 OHL CONDUCTOR LV_recover", CV7_OHL_CONDUCTOR_LV_recover, "Length (m)"),
    ("CV7 SWITCHGEAR", CV7_SWITCHGEAR, "Quantity"),
    ("CV7_UG", CV7_UG, "Quantity"),
    ("CV7_CB", CV7_CB, "Quantity"),
    ("Switch", Switch, "Quantity"),
    ("Fuses", Fuses, "Quantity"),
    ("CV31", CV31, "Quantity"),
]

def sanitize_sheet_name(name: str) -> str:
    name = str(name)
    name = re.sub(r'[:\\/*?\[\]\n\r]', '_', name)
    name = re.sub(r'[^\x00-\x7F]', '_', name)
    return name[:31]

erect_h_items = [k for k in CV7_erect.keys() if "'H' HV/EHV Pole" in k]
recover_h_items = [k for k in CV7_recover.keys() if "'A' / 'H' pole" in k]
# -------------------
# DATA COLLECTION DICTS
# -------------------
bar_data_dict = {}       # Will hold bar chart data per category
drilldown_dict = {}      # Will hold drill-down tables per category

for cat_name, keys, y_label in categories:

    if 'item' not in filtered_df.columns or 'mapped' not in filtered_df.columns:
        st.warning("Missing required columns: item / mapped")
        continue

    pattern = '|'.join([re.escape(k) for k in keys.keys()])
    mask = filtered_df['item'].astype(str).str.contains(pattern, case=False, na=False)
    sub_df = filtered_df[mask]

    # --- Normalize dates in sub_df ---
    for col in ['datetouse_dt', 'plan1_display', 'done_display']:
        if col in sub_df.columns:
            sub_df[col + "_display"] = pd.to_datetime(sub_df[col], errors='coerce') \
                .dt.strftime("%d/%m/%Y") \
                .fillna("Missing")

    # --- Clean numeric columns ---
    sub_df['qcvi_clean'] = pd.to_numeric(sub_df['qcvi'] if 'qcvi' in sub_df.columns else pd.Series(0, index=sub_df.index), errors='coerce').fillna(0)
    sub_df['qsub_clean'] = pd.to_numeric(sub_df['qsub'] if 'qsub' in sub_df.columns else pd.Series(0, index=sub_df.index), errors='coerce').fillna(0)
    sub_df["multiplier"] = 1
    sub_df.loc[sub_df["item"].isin(erect_h_items), "multiplier"] = 2
    sub_df.loc[sub_df["item"].isin(recover_h_items), "multiplier"] = 2
    sub_df["adj_value"] = sub_df["qsub_clean"] * sub_df["multiplier"]


    if cat_name == "CV31":
        # --- Collect CV7 items ---
        cv7_items = set().union(
            *[cat.keys() for cat in [CV7_erect, CV7_erect_H, CV7_erect_lv, CV7_recover] if cat]
        )

        # --- Get poles used in CV7 ---
        cv7_poles = sub_df.loc[
            sub_df['item'].isin(cv7_items), 'pole'
        ].dropna().unique()

        # --- Exclude CV7 poles from CV31 ---
        cv31_filtered = sub_df.loc[
            (~sub_df['pole'].isin(cv7_poles)) &
            (sub_df['pole'].notna())
        ].copy()
        
        cv31_unique = (
            cv31_filtered
            .sort_values(by='mapped')  # optional, controls which category "wins"
            .drop_duplicates(subset='pole', keep='first')
        )

        # --- Aggregate ---
        bar_data = cv31_unique.groupby('mapped').agg(
            Total=('pole', 'count'),
            Variation=('qcvi_clean', 'sum')
        ).reset_index()

        # --- Drilldown matches filtered unique poles ---
        drilldown_dict[cat_name] = cv31_unique.copy()

    else:
        bar_data = sub_df.groupby('mapped').agg(
            Total=('adj_value', 'sum'),
            Variation=('qcvi_clean', 'sum')
        ).reset_index()

        if cat_name != "CV31":
            drilldown_dict[cat_name] = sub_df.copy()

    bar_data.rename(columns={'mapped':'Mapped'}, inplace=True)
    bar_data['PositiveVar'] = bar_data['Variation'].clip(lower=0)
    bar_data['NegativeVar'] = bar_data['Variation'].clip(upper=0)

    # Convert to miles if needed
    y_axis_label = y_label
    if cat_name in ["CV7 OHL CONDUCTOR_instal","CV7 OHL CONDUCTOR LV_instal"] and convert_to_miles:
        bar_data['Total'] = bar_data['Total'] * 0.621371
        y_axis_label = "Length (Miles)"

    grand_total = bar_data['Total'].sum()

    # Add to bar data dict
    bar_data_dict[cat_name] = bar_data

    # Add to drill-down dict
    if cat_name != "CV31":
        drilldown_dict[cat_name] = sub_df.copy()
    st.subheader(f"🔹 {cat_name} — Total: {grand_total:,.2f}")

    # --- Plot bar chart only if there is data ---
    if grand_total > 0:
        fig = go.Figure()
        fig.add_bar(
            x=bar_data['Mapped'], y=bar_data['Total'],
            name="Quantity", marker_color="#4C78A8", text=bar_data['Total'],
            texttemplate='%{y:,.1f}', textposition='outside'
        )
        fig.add_bar(
            x=bar_data['Mapped'], y=bar_data['PositiveVar'],
            name="Positive Variation", marker_color="green"
        )
        fig.add_bar(
            x=bar_data['Mapped'], y=bar_data['NegativeVar'],
            name="Negative Variation", marker_color="red"
        )
        fig.update_layout(
            barmode='relative',
            title=f"{cat_name} Overview",
            xaxis_title="Mapping",
            yaxis_title=y_axis_label,
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            yaxis=dict(gridcolor='rgba(255,255,255,0.3)')
        )
        st.plotly_chart(fig, use_container_width=True, height=500)
    else:
        st.info(f"No data available for {cat_name}, chart not displayed.")

    # --- Collapsible drill-down ---
    with st.expander("🔍 Click to explore more information", expanded=False):
        st.subheader("Select Mapping to Drill-down:")
        cols = st.columns(3)
        for idx, mapping_value in enumerate(bar_data['Mapped']):
            col_idx = idx % 3
            with cols[col_idx]:
                button_key = f"btn_{cat_name}_{mapping_value}_{idx}"
                if st.button(f"📊 {mapping_value}", key=button_key, use_container_width=True):
                    st.session_state[f"selected_{cat_name}"] = mapping_value
                    st.rerun()

        selected_mapping = st.session_state.get(f"selected_{cat_name}")
        if selected_mapping:
            st.subheader(f"Details for: **{selected_mapping}**")
            if st.button("❌ Clear Selection", key=f"clear_{cat_name}"):
                del st.session_state[f"selected_{cat_name}"]
                st.rerun()

            selected_rows = drilldown_dict[cat_name]
            selected_rows = selected_rows[selected_rows['mapped'] == selected_mapping].copy()
            selected_rows.columns = selected_rows.columns.str.strip().str.lower()
            display_columns = [
                'shire', 'project', 'segmentcode', 'segmentdesc', 'comment',
                'pole', 'qty', 'qcvi', 'qsub', 'plan1', 'done', 'item'
            ]
            display_columns = [c for c in display_columns if c in selected_rows.columns]
            display_df = selected_rows[display_columns].copy()
            display_df.rename(columns={
                'shire': 'District',
                'segmentcode':'Circuit',
                'segmentdesc': 'Segment',
                'qty': 'Quantity',
                'qcvi':'Variation',
                'qsub': 'Quantity Used'
            }, inplace=True)
            st.write(f"**Total records:** {len(display_df)}")

            # Display table
            st.write("🔹 Information Resumed:")
            st.dataframe(display_df, use_container_width=True)


# --------------------------------------------------
# CV8 CALCULATION (EXCLUDE CV7 POLES)
# --------------------------------------------------

# --------------------------------------------------
# MAIN CV8 FUNCTION
# --------------------------------------------------
def run_cv8_analysis(filtered_df, CV7_erect, CV7_erect_H, CV7_erect_lv, CV7_recover, CV8):
    # SAFETY CHECK
    if filtered_df is None or filtered_df.empty:
        st.error("Data not loaded into filtered_df")
        st.stop()

    # -------------------------------
    # BAR CHART FUNCTION
    # -------------------------------
    def plot_bar_chart(df, category_name, x_col, y_col="Total", y_label="Quantity"):
        if df.empty:
            st.warning(f"No data for {category_name}")
            return

        df = df.sort_values(by=y_col, ascending=False)
        fig = go.Figure()
        fig.add_bar(
            x=df[x_col],
            y=df[y_col],
            text=df[y_col],
            textposition='outside',
            marker_color="#4C78A8"
        )
        fig.update_layout(
            title=f"{category_name} Overview",
            xaxis_title=x_col,
            yaxis_title=y_label,
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            yaxis=dict(gridcolor='rgba(255,255,255,0.3)'),
            xaxis_tickangle=-30
        )
        st.plotly_chart(fig, use_container_width=True)

    # -------------------------------
    # COLLECT CV7 POLES
    # -------------------------------
    cv7_items = set().union(*[
        cat.keys() for cat in [CV7_erect, CV7_erect_H ,CV7_erect_lv, CV7_recover] if cat
    ])
    cv7_poles = filtered_df.loc[filtered_df['item'].isin(cv7_items), 'pole'].dropna().unique()

    # FILTER CV8 POLES
    CV8_items = set(CV8.keys())
    cv8_df = filtered_df.loc[
        (~filtered_df['pole'].isin(cv7_poles)) &
        (filtered_df['pole'].notna()) &
        (filtered_df['item'].isin(CV8_items))
    ].copy()

    # TYPE ASSIGNMENT
    cv8_df['CV8_type'] = np.where(
        cv8_df['project'].astype(str).str.upper().str.contains('LV', na=False),
        'CV8_LV',
        'CV8_HV'
    )

    # AGGREGATE SUMMARY
    cv8_summary = cv8_df.groupby('CV8_type', as_index=False)['pole'].nunique().rename(columns={'pole':'Total'})

    # PLOT BAR CHART
    plot_bar_chart(cv8_summary, "CV8 Unique Poles", x_col="CV8_type", y_col="Total", y_label="Unique Poles")

    # -------------------------------
    # DATE NORMALISATION
    # -------------------------------
    date_cols = ['datetouse_dt', 'plan1', 'done']
    existing_cols = [col for col in date_cols if col in cv8_df.columns]

    if existing_cols:
        formatted_dates = (
            cv8_df[existing_cols]
            .apply(pd.to_datetime, errors='coerce')
            .apply(lambda col: col.dt.strftime("%d/%m/%Y"))
            .fillna("Missing")
        )
        formatted_dates.columns = [col + '_display' for col in existing_cols]
        cv8_df = pd.concat([cv8_df, formatted_dates], axis=1)

    for col in set(date_cols) - set(existing_cols):
        cv8_df[col + '_display'] = "Missing"

    # -------------------------------
    # DRILL-DOWN TABLE
    # -------------------------------
    with st.expander("🔍 CV8 Drill-down: Unique Poles Details", expanded=False):
        display_cols = [
            'project', 'segmentcode', 'segmentdesc',
            'pole', 'item', 'comment',
            'plan1_display', 'done_display'
        ]
        display_cols = [c for c in display_cols if c in cv8_df.columns]

        display_df = cv8_df[display_cols].drop_duplicates(subset='pole').sort_values('pole')
        st.dataframe(display_df, use_container_width=True)
        st.write(f"**Total unique poles displayed:** {display_df['pole'].nunique()}")

    return cv8_df, cv8_summary

# --------------------------------------------------
# CALL CV8 FUNCTION
# --------------------------------------------------
cv8_df, cv8_summary = run_cv8_analysis(
    filtered_df,
    CV7_erect,
    CV7_erect_H,
    CV7_erect_lv,
    CV7_recover,
    CV8
)

# --------------------------------------------------
# PAGE LAYOUT
# --------------------------------------------------
st.set_page_config(layout="wide")
st.markdown("""
    <style>
        .block-container { padding-top:1rem; padding-bottom:1rem; max-width:100%; }
        .scroll-box {
            max-height:400px; overflow-y:auto; overflow-x:auto;
            padding:12px; border:1px solid #444; background-color:#111;
            font-family:monospace; font-size:14px; white-space:nowrap; color:#fff;
        }
        .scroll-box span { display:inline-block; min-width:120px; padding-right:20px; }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h2 style='text-align:center; color:white;'>Projects & Circuits Overview</h2>", unsafe_allow_html=True)

required_cols = ['shire', 'datetouse_dt', 'project', 'segmentcode', 'segmentdesc']
existing_cols = [c for c in required_cols if c in filtered_df.columns]

if 'project' in existing_cols:
    projects = filtered_df['project'].dropna().unique()

    for proj in sorted(projects):
        proj_df = filtered_df[filtered_df['project'] == proj]
        cols_to_use = [c for c in required_cols if c in proj_df.columns]
        segments = proj_df[cols_to_use].dropna(subset=['segmentcode']).drop_duplicates()

        with st.expander(f"Project: {proj} ({len(segments)} circuits)"):
            display_lines = []
            for _, row in segments.iterrows():
                district = str(row.get("shire", ""))
                dt = row.get("datetouse_dt", None)
                date = dt.strftime("%d/%m/%Y") if pd.notna(dt) else "Missing"
                circuit = str(row.get("segmentcode", ""))
                segment = str(row.get("segmentdesc", ""))
                display_lines.append(f"{district} | {date} | {circuit} | {segment}")

            st.markdown(
                "<div class='scroll-box'>" + "<br>".join(display_lines) + "</div>",
                unsafe_allow_html=True
            )
else:
    st.info("Project column not found in the data.")

# -------------------------------
# HIGH LEVEL PLANNING EXPORT BUTTON
# -------------------------------
col1, center_col, col3 = st.columns([1, 3, 1])
with center_col:
    if 'filtered_df' in locals() and not filtered_df.empty:
        # Columns for the high level planning
        export_columns = [
            "shire",
            "datetouse_display",
            "project",
            "segmentcode",
            "segmentdesc"
        ]
        export_df = filtered_df[[c for c in export_columns if c in filtered_df.columns]].copy()

        excel_file_high_level = generate_excel_styled_multilevel(
            export_df,
            poles_df if 'poles_df' in locals() else None
        )

        st.download_button(
            label="📥 High Level Planning & Poles Excel",
            data=excel_file_high_level,
            file_name=f"High_level_planning_{date_range_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


# --------------------------------------------------
# EXCEL EXPORT
# --------------------------------------------------
def sanitize_sheet_name(name: str) -> str:
    name = str(name)
    name = re.sub(r'[:\\/*?\[\]\n\r]', '_', name)
    name = re.sub(r'[^\x00-\x7F]', '_', name)
    return name[:31]

display_columns = [
    'shire', 'project', 'segmentcode', 'segmentdesc', 'comment',
    'pole', 'qty', 'qcvi', 'qsub', 'plan1', 'done', 'item','total','orig'
]
def generate_excel_export(display_columns, drilldown_dict, cv8_df, filtered_df):
    output = io.BytesIO()

    # -----------------------------
    # Prepare individual sheets
    # -----------------------------
    def prepare_df(df):
        df = df.copy()
        for col in display_columns:
            if col not in df.columns:
                df[col] = ""
        return df[display_columns].fillna("")

    all_data = {}
    for name, df in drilldown_dict.items():
        if not df.empty:
            all_data[name] = prepare_df(df)
    if cv8_df is not None and not cv8_df.empty:
        all_data["CV8"] = prepare_df(cv8_df)

    # -----------------------------
    # Combined_Data: all filtered info
    # -----------------------------
    combined_df = filtered_df.copy()

    # -----------------------------
    # Build Project Summary
    # -----------------------------
    summary_rows = []
    if not combined_df.empty:
        all_projects = combined_df['project'].dropna().unique()

        for project in all_projects:
            row = {"project": project}

            # Per-category values for chart purposes
            for name, df in all_data.items():
                proj_df = df[df['project'] == project]
                if name in ["CV31", "CV8"]:
                    val = proj_df['pole'].nunique() if 'pole' in proj_df.columns else 0
                else:
                    val = pd.to_numeric(proj_df['qsub'], errors='coerce').fillna(0).sum() \
                        if 'qsub' in proj_df.columns else 0
                row[name] = val

            # Sum Total & Original from Combined_Data
            proj_combined = combined_df[combined_df['project'] == project]
            row["total"] = proj_combined['total'].sum() if 'total' in proj_combined.columns else 0
            row["orig"] = proj_combined['orig'].sum() if 'orig' in proj_combined.columns else 0

            summary_rows.append(row)

        summary_df = pd.DataFrame(summary_rows)
    else:
        summary_df = pd.DataFrame()

    # -----------------------------
    # Openpyxl Workbook
    # -----------------------------
    wb = Workbook()
    wb.remove(wb.active)  # remove default sheet

    # Formatting
    IMG_HEIGHT = 120
    IMG_WIDTH_SMALL = 120
    IMG_WIDTH_LARGE = IMG_WIDTH_SMALL * 3

    header_font = Font(bold=True, size=16)
    header_fill = PatternFill(start_color="00CCFF", end_color="00CCFF", fill_type="solid")
    thin_side = Side(style="thin")
    medium_side = Side(style="medium")
    thick_side = Side(style="thick")
    light_grey_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    red_font = Font(color="FF0000")
    green_font = Font(color="00AA00")
    black_font = Font(color="000000")

    # -----------------------------
    # 1️⃣ Project Summary sheet
    # -----------------------------
    if not summary_df.empty:
        ws = wb.create_sheet("Project_Summary")
        ws.append(summary_df.columns.tolist())
        for idx, row in summary_df.iterrows():
            ws.append(row.tolist())

        # Header formatting
        for col_idx, cell in enumerate(ws[1], start=1):
            cell.font = header_font
            cell.fill = header_fill
            ws.column_dimensions[get_column_letter(col_idx)].width = 20
            cell.border = Border(
                left=thick_side if col_idx == 1 else medium_side,
                right=thick_side if col_idx == ws.max_column else medium_side,
                top=thick_side,
                bottom=thick_side
            )

        # Row alternating colors
        for row_idx in range(2, ws.max_row + 1):
            fill = light_grey_fill if row_idx % 2 == 0 else white_fill
            for col_idx in range(1, ws.max_column + 1):
                ws.cell(row=row_idx, column=col_idx).fill = fill


    # -----------------------------
    # 2️⃣ Individual Sheets
    # -----------------------------
    for name, df in all_data.items():
        ws = wb.create_sheet(sanitize_sheet_name(name))
        ws.append(df.columns.tolist())
        for _, row in df.iterrows():
            ws.append(row.tolist())

        # Header formatting
        for col_idx, cell in enumerate(ws[1], start=1):
            cell.font = header_font
            cell.fill = header_fill
            ws.column_dimensions[get_column_letter(col_idx)].width = 20
            cell.border = Border(
                left=thick_side if col_idx == 1 else medium_side,
                right=thick_side if col_idx == ws.max_column else medium_side,
                top=thick_side,
                bottom=thick_side
            )

        # Row alternating colors
        for row_idx in range(2, ws.max_row + 1):
            fill = light_grey_fill if row_idx % 2 == 0 else white_fill
            for col_idx in range(1, ws.max_column + 1):
                ws.cell(row=row_idx, column=col_idx).fill = fill

    # -----------------------------
    # 3️⃣ Combined_Data Sheet
    # -----------------------------
    if not combined_df.empty:
        ws = wb.create_sheet("Combined_Data")
        ws.append(combined_df.columns.tolist())
        for _, row in combined_df.iterrows():
            ws.append(row.tolist())

        # Header formatting
        for col_idx, cell in enumerate(ws[1], start=1):
            cell.font = header_font
            cell.fill = header_fill
            ws.column_dimensions[get_column_letter(col_idx)].width = 20
            cell.border = Border(
                left=thick_side if col_idx == 1 else medium_side,
                right=thick_side if col_idx == ws.max_column else medium_side,
                top=thick_side,
                bottom=thick_side
            )

        # Row alternating colors
        for row_idx in range(2, ws.max_row + 1):
            fill = light_grey_fill if row_idx % 2 == 0 else white_fill
            for col_idx in range(1, ws.max_column + 1):
                ws.cell(row=row_idx, column=col_idx).fill = fill

    # Save to BytesIO
    wb.save(output)
    output.seek(0)
    return output.getvalue()


# -------------------------------
# DOWNLOAD BUTTON
# -------------------------------
if drilldown_dict or (cv8_df is not None and not cv8_df.empty):
    excel_bytes = generate_excel_export(display_columns, drilldown_dict, cv8_df, filtered_df)
    st.download_button(
        label="📥 Export All Data to Excel",
        data=excel_bytes,
        file_name=f"Planning_Export_{date_range_str}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
